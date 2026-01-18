"""
Document loaders for ingestion pipeline.
Supports Markdown, Text, PDF, Word (.docx), Excel (.xlsx), and RTF files.

Can load from file paths or from bytes (for uploaded files).
"""
from io import BytesIO
from pathlib import Path
from typing import Iterator

try:
    from pypdf import PdfReader
except ImportError:
    PdfReader = None

try:
    from docx import Document as DocxDocument
except ImportError:
    DocxDocument = None

try:
    from openpyxl import load_workbook
except ImportError:
    load_workbook = None

try:
    from striprtf.striprtf import rtf_to_text
except ImportError:
    rtf_to_text = None


def load_text(path: Path) -> str:
    """Load plain text or markdown file."""
    return path.read_text(encoding="utf-8")


def load_pdf(path: Path) -> str:
    """Load PDF and extract text from all pages."""
    if PdfReader is None:
        raise ImportError("pypdf is required for PDF loading. Install with: pip install pypdf")
    
    reader = PdfReader(path)
    texts = []
    for page in reader.pages:
        text = page.extract_text()
        if text:
            texts.append(text)
    return "\n\n".join(texts)


def load_docx(path: Path) -> str:
    """
    Load Word document and extract plain paragraph text in reading order.
    Ignores formatting, styles, headers, footers, and comments.
    """
    if DocxDocument is None:
        raise ImportError("python-docx is required for Word loading. Install with: pip install python-docx")
    
    doc = DocxDocument(path)
    texts = []
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:
            texts.append(text)
    return "\n\n".join(texts)


def load_xlsx(path: Path) -> tuple[str, list[dict]]:
    """
    Load Excel workbook and extract text from all sheets.
    
    Returns:
        Tuple of (full text content, list of sheet metadata dicts)
        Each sheet's content is prefixed with the sheet name.
    """
    if load_workbook is None:
        raise ImportError("openpyxl is required for Excel loading. Install with: pip install openpyxl")
    
    wb = load_workbook(path, data_only=True)  # data_only=True to get computed values, not formulas
    all_texts = []
    sheet_metadata = []
    
    for sheet_name in wb.sheetnames:
        sheet = wb[sheet_name]
        sheet_texts = []
        
        for row in sheet.iter_rows():
            cell_values = []
            for cell in row:
                if cell.value is not None:
                    cell_values.append(str(cell.value))
            
            if cell_values:  # Skip empty rows
                row_text = "\t".join(cell_values)
                sheet_texts.append(row_text)
        
        if sheet_texts:
            sheet_content = f"[Sheet: {sheet_name}]\n" + "\n".join(sheet_texts)
            all_texts.append(sheet_content)
            sheet_metadata.append({"sheet_name": sheet_name})
    
    wb.close()
    return "\n\n".join(all_texts), sheet_metadata


def load_rtf(path: Path) -> str:
    """Load RTF file and convert to plain text."""
    if rtf_to_text is None:
        raise ImportError("striprtf is required for RTF loading. Install with: pip install striprtf")

    # Read as bytes first, then try different encodings
    raw_bytes = path.read_bytes()

    # Try common encodings
    for encoding in ("utf-8", "latin-1", "cp1252"):
        try:
            rtf_content = raw_bytes.decode(encoding)
            break
        except UnicodeDecodeError:
            continue
    else:
        # Fallback to latin-1 which accepts any byte
        rtf_content = raw_bytes.decode("latin-1")

    text = rtf_to_text(rtf_content)
    return text.strip()


def load_document(path: Path) -> str | tuple[str, list[dict]]:
    """
    Load a document based on its file extension.

    Supported formats:
    - .txt, .md: Plain text
    - .pdf: PDF (requires pypdf)
    - .docx: Word (requires python-docx)
    - .xlsx: Excel (requires openpyxl) - returns tuple with sheet metadata
    - .rtf: Rich Text Format (requires striprtf)

    Returns:
        str for most formats, or tuple[str, list[dict]] for xlsx with sheet metadata
    """
    suffix = path.suffix.lower()

    if suffix in (".txt", ".md"):
        return load_text(path)
    elif suffix == ".pdf":
        return load_pdf(path)
    elif suffix == ".docx":
        return load_docx(path)
    elif suffix == ".xlsx":
        return load_xlsx(path)
    elif suffix == ".rtf":
        return load_rtf(path)
    else:
        raise ValueError(f"Unsupported file type: {suffix}")


# =============================================================================
# Loaders from bytes (for uploaded files)
# =============================================================================

def load_text_from_bytes(data: bytes) -> str:
    """Load plain text or markdown from bytes."""
    return data.decode("utf-8")


def load_pdf_from_bytes(data: bytes) -> str:
    """Load PDF from bytes and extract text from all pages."""
    if PdfReader is None:
        raise ImportError("pypdf is required for PDF loading. Install with: pip install pypdf")

    reader = PdfReader(BytesIO(data))
    texts = []
    for page in reader.pages:
        text = page.extract_text()
        if text:
            texts.append(text)
    return "\n\n".join(texts)


def load_docx_from_bytes(data: bytes) -> str:
    """Load Word document from bytes and extract plain paragraph text."""
    if DocxDocument is None:
        raise ImportError("python-docx is required for Word loading. Install with: pip install python-docx")

    doc = DocxDocument(BytesIO(data))
    texts = []
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:
            texts.append(text)
    return "\n\n".join(texts)


def load_xlsx_from_bytes(data: bytes) -> tuple[str, list[dict]]:
    """Load Excel workbook from bytes and extract text from all sheets."""
    if load_workbook is None:
        raise ImportError("openpyxl is required for Excel loading. Install with: pip install openpyxl")

    wb = load_workbook(BytesIO(data), data_only=True)
    all_texts = []
    sheet_metadata = []

    for sheet_name in wb.sheetnames:
        sheet = wb[sheet_name]
        sheet_texts = []

        for row in sheet.iter_rows():
            cell_values = []
            for cell in row:
                if cell.value is not None:
                    cell_values.append(str(cell.value))

            if cell_values:
                row_text = "\t".join(cell_values)
                sheet_texts.append(row_text)

        if sheet_texts:
            sheet_content = f"[Sheet: {sheet_name}]\n" + "\n".join(sheet_texts)
            all_texts.append(sheet_content)
            sheet_metadata.append({"sheet_name": sheet_name})

    wb.close()
    return "\n\n".join(all_texts), sheet_metadata


def load_rtf_from_bytes(data: bytes) -> str:
    """Load RTF from bytes and convert to plain text."""
    if rtf_to_text is None:
        raise ImportError("striprtf is required for RTF loading. Install with: pip install striprtf")

    # Try common encodings
    for encoding in ("utf-8", "latin-1", "cp1252"):
        try:
            rtf_content = data.decode(encoding)
            break
        except UnicodeDecodeError:
            continue
    else:
        # Fallback to latin-1 which accepts any byte
        rtf_content = data.decode("latin-1")

    text = rtf_to_text(rtf_content)
    return text.strip()


def load_document_from_bytes(data: bytes, filename: str) -> str | tuple[str, list[dict]]:
    """
    Load a document from bytes based on the filename extension.

    Args:
        data: File content as bytes.
        filename: Original filename (used to determine file type).

    Returns:
        str for most formats, or tuple[str, list[dict]] for xlsx with sheet metadata.
    """
    suffix = Path(filename).suffix.lower()

    if suffix in (".txt", ".md"):
        return load_text_from_bytes(data)
    elif suffix == ".pdf":
        return load_pdf_from_bytes(data)
    elif suffix == ".docx":
        return load_docx_from_bytes(data)
    elif suffix == ".xlsx":
        return load_xlsx_from_bytes(data)
    elif suffix == ".rtf":
        return load_rtf_from_bytes(data)
    else:
        raise ValueError(f"Unsupported file type: {suffix}")


def iter_documents(
    directory: Path,
    extensions: tuple[str, ...] = (".md", ".txt", ".pdf", ".docx", ".xlsx", ".rtf")
) -> Iterator[tuple[Path, str | tuple[str, list[dict]]]]:
    """
    Iterate over all documents in a directory.
    
    Yields:
        Tuples of (path, content) for each document.
        For xlsx files, content is a tuple of (text, sheet_metadata).
    """
    for ext in extensions:
        for path in directory.glob(f"*{ext}"):
            if path.is_file():
                try:
                    content = load_document(path)
                    yield path, content
                except ValueError as e:
                    print(f"Skipping unsupported file {path}: {e}")
                except Exception as e:
                    print(f"Error loading {path}: {e}")

